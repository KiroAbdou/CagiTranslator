import cv2
import pytesseract
import arabic_reshaper
from bidi.algorithm import get_display
from pdf2image import convert_from_path
from PIL import Image
import numpy as np
from docx import Document
import easyocr

reader = easyocr.Reader(['ar'])

unproc_dir = "./docs/unprocessed/"
process_dir = "./docs/processed/"

# Set input and output file paths
input_doc = "curriculum_5_P1_2019"
input_pdf = unproc_dir + input_doc + ".pdf"
output_pdf = process_dir + input_doc +  "_lec1.pdf"

# Manually set the Tesseract executable path
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

page = 22
# Convert page 22 to image
images = convert_from_path(input_pdf, first_page=page, last_page=page)
image = images[0]  # Get the first (and only) extracted image

# Save the extracted image
image_path = f"page_{page}.png"
# cv2.imwrite(image_path, np.array(image))
image = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)

# ====== STEP 2: Apply Adaptive Thresholding for Better Contrast ======
binary_image = cv2.adaptiveThreshold(
    image, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 15, 10
)

# ====== STEP 3: Remove Small Gaps (Morphological Closing) ======
kernel = np.ones((1, 2), np.uint8)  # Small kernel to connect characters
closed_image = cv2.morphologyEx(binary_image, cv2.MORPH_CLOSE, kernel, iterations=2)

# ====== STEP 4: Remove Small White Noise ======
kernel_noise = np.ones((2, 2), np.uint8)  # Slightly larger kernel for small artifacts
cleaned_image = cv2.morphologyEx(closed_image, cv2.MORPH_OPEN, kernel_noise, iterations=1)

# Save the processed image
processed_image_path = f"cleaned_page_{page}.png"
cv2.imwrite(processed_image_path, cleaned_image)
result_ocr = reader.readtext(processed_image_path)
print("result_ocr: \n", result_ocr)
ocr_text = "\n".join([entry[1] for entry in result_ocr])

# ====== STEP 5: Perform OCR on the Cleaned Image ======
extracted_text = pytesseract.image_to_string(Image.open(processed_image_path), lang="ara")

# ====== STEP 6: Fix Arabic Text Formatting ======
reshaped_text = arabic_reshaper.reshape(extracted_text)
corrected_text = get_display(reshaped_text)
# print("corrected_text: \n", corrected_text)

output_docx_file = f"extracted_text_P_{page}.docx"

# Create a Word document and set text alignment to RTL
doc = Document()
paragraph = doc.add_paragraph()
paragraph.alignment = 2  # Right-align text
run = paragraph.add_run(reshaped_text+"\n\n\n\n\n\n\n\nOCR\n"+ ocr_text)
run.font.rtl = True  # Set RTL text

# Save the document
doc.save(output_docx_file)

print(f"✅ Arabic text saved with RTL formatting: {output_docx_file}")

